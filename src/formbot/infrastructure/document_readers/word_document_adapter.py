from __future__ import annotations

import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Literal

from formbot.domain.exceptions import (
    DocumentProcessingError,
    DocumentSaveError,
    LabelNotFoundError,
    MappingRuleError,
    PositionOutOfBoundsError,
)
from formbot.domain.models import CellPosition
from formbot.domain.ports.document_adapter import DocumentAdapter
from formbot.shared.utils import normalize_text

LOGGER = logging.getLogger(__name__)

# Prefijos usados en CellPosition.sheet_name para identificar el tipo de destino.
# write_value() decodifica estos prefijos para enrutar la escritura al lugar correcto.
_WORD_TABLE_PREFIX = "word_table_"
_WORD_PARA_SHEET = "word_paragraph"

try:
    from docx import Document as DocxDocument

    _DOCX_AVAILABLE = True
except ImportError:  # pragma: no cover
    _DOCX_AVAILABLE = False


@dataclass(frozen=True)
class _WordLabelCandidate:
    """Candidato de etiqueta encontrado en el documento Word."""

    position: CellPosition
    match_type: Literal["exact", "partial"]
    text: str  # Texto original de la celda / párrafo


class WordDocumentAdapter(DocumentAdapter):
    """Adaptador para formularios Word (.docx) usando python-docx.

    Estrategia de búsqueda de etiquetas:
    - Busca primero en celdas de tablas (caso más común en formularios empresariales).
    - Si no encuentra, busca en párrafos del cuerpo del documento.
    - Normaliza el texto eliminando colons, acentos y diferencias de mayúsculas.
    - Soporta coincidencia exacta y parcial, igual que el adaptador Excel.

    Codificación de CellPosition:
    - Celda de tabla:   sheet_name="word_table_N", row=fila+1, column=columna+1
    - Párrafo:          sheet_name="word_paragraph", row=párrafo+1, column=1

    Esto permite que el LabelOffsetMapper aplique offsets directamente:
    - column_offset=1 en una tabla → columna adyacente de la misma fila (patrón típico).
    - row_offset=1 en párrafos    → siguiente párrafo.

    Estrategia de escritura:
    - Para celdas: reemplaza el contenido de la celda destino preservando el estilo del primer run.
    - Para párrafos: reemplaza el contenido del párrafo destino.

    Limitaciones conocidas:
    - El modo precisión (PrecisionFillUseCase) es Excel-exclusivo — usar FillFormUseCase.
    - No detecta Content Controls (SDTs) en esta versión.
    - El parámetro sheet_name de find_label() se ignora (Word no tiene hojas).
    - No maneja celdas combinadas (merged cells) de tabla.
    """

    def __init__(self, template_path: Path) -> None:
        if not _DOCX_AVAILABLE:
            raise DocumentProcessingError(
                "La librería 'python-docx' no está instalada. "
                "Instálela con:  pip install python-docx>=1.1.2"
            )
        if not template_path.exists():
            raise DocumentProcessingError(
                f"No existe el template Word: {template_path}"
            )
        if template_path.suffix.lower() not in {".docx"}:
            raise DocumentProcessingError(
                f"Se esperaba un archivo .docx, se recibió '{template_path.suffix}': {template_path}"
            )
        try:
            self._document: Any = DocxDocument(str(template_path))
        except Exception as exc:
            raise DocumentProcessingError(
                f"No fue posible abrir el documento Word: {template_path}"
            ) from exc

        self._template_path = template_path
        LOGGER.debug(
            "Word '%s' cargado: %d tabla(s), %d párrafo(s).",
            template_path.name,
            len(self._document.tables),
            len(self._document.paragraphs),
        )

    # ------------------------------------------------------------------
    # Interfaz pública (DocumentAdapter)
    # ------------------------------------------------------------------

    def find_label(self, text: str, sheet_name: str | None = None) -> CellPosition:
        if sheet_name is not None:
            LOGGER.warning(
                "WordDocumentAdapter: sheet_name='%s' ignorado — "
                "Word no tiene hojas. Se busca en todo el documento.",
                sheet_name,
            )

        normalized_target = normalize_text(text)
        if not normalized_target:
            raise LabelNotFoundError("No se puede buscar una etiqueta vacía")

        exact_table, partial_table = self._collect_table_candidates(normalized_target)
        exact_para, partial_para = self._collect_para_candidates(normalized_target)

        # Las tablas tienen prioridad sobre los párrafos
        all_exact = exact_table + exact_para
        all_partial = partial_table + partial_para

        if len(all_exact) == 1:
            chosen = all_exact[0]
            LOGGER.debug(
                "Etiqueta '%s' encontrada (exact) en Word: %s @ %s",
                text,
                chosen.text,
                chosen.position,
            )
            return chosen.position

        if len(all_exact) > 1:
            locations = " | ".join(
                f"{c.text!r} @ {c.position.sheet_name} R{c.position.row}C{c.position.column}"
                for c in all_exact
            )
            raise MappingRuleError(
                f"Etiqueta ambigua '{text}': {len(all_exact)} coincidencias exactas "
                f"en el documento Word ({locations})"
            )

        if len(all_partial) == 1:
            chosen = all_partial[0]
            LOGGER.debug(
                "Etiqueta '%s' encontrada (partial) en Word: %s @ %s",
                text,
                chosen.text,
                chosen.position,
            )
            return chosen.position

        if len(all_partial) > 1:
            locations = " | ".join(
                f"{c.text!r} @ {c.position.sheet_name} R{c.position.row}C{c.position.column}"
                for c in all_partial
            )
            raise MappingRuleError(
                f"Etiqueta ambigua '{text}': {len(all_partial)} coincidencias parciales "
                f"en el documento Word ({locations})"
            )

        raise LabelNotFoundError(
            f"No se encontró etiqueta '{text}' en el documento Word "
            f"({len(self._document.tables)} tabla(s), {len(self._document.paragraphs)} párrafo(s))"
        )

    def write_value(self, position: CellPosition, value: Any) -> None:
        str_value = "" if value is None else str(value)

        if position.sheet_name.startswith(_WORD_TABLE_PREFIX):
            self._write_to_table_cell(position, str_value)
        elif position.sheet_name == _WORD_PARA_SHEET:
            self._write_to_paragraph(position, str_value)
        else:
            raise PositionOutOfBoundsError(
                f"Posición Word no reconocida: '{position.sheet_name}'. "
                f"Se esperaba '{_WORD_TABLE_PREFIX}N' o '{_WORD_PARA_SHEET}'"
            )

    def save(self, output_path: Path) -> None:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        try:
            self._document.save(str(output_path))
        except OSError as exc:
            raise DocumentSaveError(
                f"No fue posible guardar el documento Word de salida: {output_path}"
            ) from exc
        LOGGER.info("Documento Word guardado en '%s'.", output_path)

    def close(self) -> None:
        # python-docx no requiere cierre explícito
        return None

    # ------------------------------------------------------------------
    # Búsqueda de candidatos
    # ------------------------------------------------------------------

    def _collect_table_candidates(
        self,
        normalized_target: str,
    ) -> tuple[list[_WordLabelCandidate], list[_WordLabelCandidate]]:
        """Busca la etiqueta en todas las celdas de todas las tablas del documento."""
        exact: list[_WordLabelCandidate] = []
        partial: list[_WordLabelCandidate] = []

        for t_idx, table in enumerate(self._document.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    raw_text = cell.text.strip()
                    if not raw_text:
                        continue
                    norm = normalize_text(raw_text)
                    if not norm:
                        continue

                    position = CellPosition(
                        sheet_name=f"{_WORD_TABLE_PREFIX}{t_idx}",
                        row=r_idx + 1,   # 1-indexado (compatible con shifted())
                        column=c_idx + 1,
                    )

                    if norm == normalized_target:
                        exact.append(
                            _WordLabelCandidate(position, "exact", raw_text)
                        )
                    elif normalized_target in norm:
                        partial.append(
                            _WordLabelCandidate(position, "partial", raw_text)
                        )

        return exact, partial

    def _collect_para_candidates(
        self,
        normalized_target: str,
    ) -> tuple[list[_WordLabelCandidate], list[_WordLabelCandidate]]:
        """Busca la etiqueta en todos los párrafos del cuerpo del documento."""
        exact: list[_WordLabelCandidate] = []
        partial: list[_WordLabelCandidate] = []

        for p_idx, para in enumerate(self._document.paragraphs):
            raw_text = para.text.strip()
            if not raw_text:
                continue
            norm = normalize_text(raw_text)
            if not norm:
                continue

            position = CellPosition(
                sheet_name=_WORD_PARA_SHEET,
                row=p_idx + 1,  # 1-indexado
                column=1,
            )

            if norm == normalized_target:
                exact.append(_WordLabelCandidate(position, "exact", raw_text))
            elif normalized_target in norm:
                partial.append(_WordLabelCandidate(position, "partial", raw_text))

        return exact, partial

    # ------------------------------------------------------------------
    # Escritura en destinos
    # ------------------------------------------------------------------

    def _write_to_table_cell(self, position: CellPosition, value: str) -> None:
        """Escribe un valor en la celda de tabla indicada por position."""
        suffix = position.sheet_name[len(_WORD_TABLE_PREFIX):]
        try:
            t_idx = int(suffix)
        except ValueError as exc:
            raise PositionOutOfBoundsError(
                f"Índice de tabla inválido en posición Word: '{position.sheet_name}'"
            ) from exc

        if t_idx >= len(self._document.tables):
            raise PositionOutOfBoundsError(
                f"Tabla {t_idx} no existe (el documento tiene "
                f"{len(self._document.tables)} tabla(s))"
            )

        table = self._document.tables[t_idx]
        row_idx = position.row - 1    # volver a 0-indexado
        col_idx = position.column - 1

        if row_idx < 0 or row_idx >= len(table.rows):
            raise PositionOutOfBoundsError(
                f"Fila {position.row} fuera de rango en tabla {t_idx} "
                f"({len(table.rows)} fila(s))"
            )

        row_obj = table.rows[row_idx]
        if col_idx < 0 or col_idx >= len(row_obj.cells):
            raise PositionOutOfBoundsError(
                f"Columna {position.column} fuera de rango en tabla {t_idx}, "
                f"fila {position.row} ({len(row_obj.cells)} columna(s))"
            )

        cell = row_obj.cells[col_idx]
        _overwrite_cell_content(cell, value)
        LOGGER.debug(
            "Word tabla %d R%dC%d ← '%s'", t_idx, position.row, position.column, value
        )

    def _write_to_paragraph(self, position: CellPosition, value: str) -> None:
        """Escribe un valor reemplazando el contenido de un párrafo."""
        para_idx = position.row - 1  # 0-indexado
        if para_idx < 0 or para_idx >= len(self._document.paragraphs):
            raise PositionOutOfBoundsError(
                f"Párrafo {position.row} fuera de rango "
                f"({len(self._document.paragraphs)} párrafo(s))"
            )
        para = self._document.paragraphs[para_idx]
        _overwrite_para_content(para, value)
        LOGGER.debug("Word párrafo %d ← '%s'", position.row, value)


# ------------------------------------------------------------------
# Helpers de escritura (funciones puras, sin estado)
# ------------------------------------------------------------------

def _overwrite_cell_content(cell: Any, value: str) -> None:
    """Reemplaza el contenido de la primera celda de tabla, preservando el estilo."""
    if not cell.paragraphs:
        cell.add_paragraph(value)
        return

    para = cell.paragraphs[0]
    if para.runs:
        # Preservar formato del primer run; limpiar el resto
        para.runs[0].text = value
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(value)


def _overwrite_para_content(para: Any, value: str) -> None:
    """Reemplaza el contenido de un párrafo, preservando el estilo del primer run."""
    if para.runs:
        para.runs[0].text = value
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(value)
